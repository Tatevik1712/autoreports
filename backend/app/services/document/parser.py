"""
Парсинг документов: PDF, DOCX, XLSX, TXT, изображения (OCR).

Улучшения:
- PDF: используется pymupdf4llm.to_markdown() — сохраняет структуру заголовков и таблиц
- DOCX: разбиение с учётом иерархии заголовков
- XLSX: исправлен баг wb.sheetnames после wb.close()
- TXT: исправлена переменная encoding в блоке for/else
"""
from __future__ import annotations
import io
from dataclasses import dataclass, field
from pathlib import Path
from app.core.logging import get_logger
from docx import Document

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    """Результат парсинга любого документа."""
    filename: str
    content_type: str
    text: str                                           # весь извлечённый текст
    pages: list[str] = field(default_factory=list)     # текст по страницам / листам
    tables: list[list[list[str]]] = field(default_factory=list)  # таблицы как 2D-массивы
    chunks: list[dict] = field(default_factory=list)   # структурные чанки с иерархией (для RAG)
    meta: dict = field(default_factory=dict)            # доп. метаданные
    error: str | None = None


class DocumentParser:
    """Единая точка входа для парсинга любого поддерживаемого формата."""

    async def parse(self, content: bytes, filename: str, content_type: str) -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        logger.info("document_parse_start", filename=filename, ext=ext, size=len(content))

        try:
            if ext == ".pdf":
                return await self._parse_pdf(content, filename, content_type)
            elif ext in (".doc", ".docx"):
                return await self._parse_docx(content, filename, content_type)
            elif ext in (".xls", ".xlsx"):
                return await self._parse_xlsx(content, filename, content_type)
            elif ext == ".txt":
                return await self._parse_txt(content, filename, content_type)
            elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
                return await self._parse_image(content, filename, content_type)
            else:
                return ParsedDocument(
                    filename=filename,
                    content_type=content_type,
                    text="",
                    error=f"Неподдерживаемый формат: {ext}",
                )
        except Exception as exc:
            logger.error("document_parse_error", filename=filename, error=str(exc))
            return ParsedDocument(
                filename=filename,
                content_type=content_type,
                text="",
                error=f"Ошибка при обработке файла: {exc}",
            )

    # PDF
    # Используем pymupdf4llm.to_markdown() вместо page.get_text().
    # Это даёт Markdown-разметку с заголовками (#, ##), таблицами и жирным —
    # RAG-pipeline потом лучше разбивает такой текст на смысловые чанки.
    async def _parse_pdf(
        self, content: bytes, filename: str, content_type: str
    ) -> ParsedDocument:
        try:
            import pymupdf                      # fitz
            from pymupdf4llm import to_markdown
        except ImportError:
            # Fallback: если pymupdf4llm не установлен — используем обычный текст
            return await self._parse_pdf_fallback(content, filename, content_type)

        doc = pymupdf.open(stream=content, filetype="pdf")
        page_count = doc.page_count

        # to_markdown конвертирует весь документ в структурированный Markdown
        full_text = to_markdown(doc)

        # Дополнительно извлекаем таблицы через нативный API
        tables: list[list[list[str]]] = []
        for page in doc:
            try:
                for table in page.find_tables():
                    extracted = table.extract()
                    if extracted:
                        tables.append(extracted)
            except Exception:
                pass

        doc.close()

        return ParsedDocument(
            filename=filename,
            content_type=content_type,
            text=full_text,
            pages=[full_text],   # PDF как единый Markdown-документ
            tables=tables,
            meta={
                "page_count": page_count,
                "table_count": len(tables),
                "format": "markdown",  # сигнал для RAG что текст структурирован
            },
        )

    async def _parse_pdf_fallback(
        self, content: bytes, filename: str, content_type: str
    ) -> ParsedDocument:
        """Fallback без pymupdf4llm — простой текст по страницам."""
        import pymupdf

        doc = pymupdf.open(stream=content, filetype="pdf")
        pages: list[str] = []
        tables: list[list[list[str]]] = []

        for page in doc:
            pages.append(page.get_text("text"))
            try:
                for table in page.find_tables():
                    extracted = table.extract()
                    if extracted:
                        tables.append(extracted)
            except Exception:
                pass

        doc.close()
        full_text = "\n\n".join(pages)

        return ParsedDocument(
            filename=filename,
            content_type=content_type,
            text=full_text,
            pages=pages,
            tables=tables,
            meta={"page_count": len(pages), "table_count": len(tables)},
        )

    # DOCX
    # Алгоритм разбиения с учётом иерархии заголовков:
    # взят из ext_module/backend/src/utils/chunk_splitter.py (split_docx).
    # Заголовки "Заголовок 1", "Заголовок 2" ... определяют структуру документа.
    # Каждый чанк получает поле "title" — путь в иерархии вида "Глава > Раздел".
    # Это значительно улучшает качество RAG: модель видит контекст секции.
    async def _parse_docx(
        self, content: bytes, filename: str, content_type: str
    ) -> ParsedDocument:

        doc = Document(io.BytesIO(content))

        chunks: list[dict] = []       # структурные чанки с иерархией
        tables: list[list[list[str]]] = []
        heading_stack: list[tuple[int, str]] = []   # (уровень, текст заголовка)
        current_paragraphs: list[str] = []

        def _flush_chunk() -> None:
            """Сохраняет накопленные параграфы как чанк."""
            text = "\n".join(p for p in current_paragraphs if p.strip())
            if text:
                hierarchy = " > ".join(h[1] for h in heading_stack)
                chunks.append({
                    "title": hierarchy,
                    "content": text,
                    "level": len(heading_stack),
                })
            current_paragraphs.clear()

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            # Определяем заголовок: "Заголовок 1", "Heading 1" и т.п.
            is_heading = (
                style_name.startswith("Заголовок")
                or style_name.startswith("Heading")
            )

            if is_heading:
                # Извлекаем уровень из имени стиля (последний символ)
                try:
                    level = int(style_name[-1])
                except ValueError:
                    level = 1

                # Сохраняем накопленный текст перед новым заголовком
                _flush_chunk()

                # Обновляем стек заголовков: убираем всё с уровнем >= текущего
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()

                heading_stack.append((level, para.text.strip()))

            else:
                # Обычный параграф — добавляем в текущий чанк
                if para.text.strip():
                    current_paragraphs.append(para.text)

        # Сохраняем последний чанк
        _flush_chunk()

        # Извлекаем таблицы
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables.append(rows)

        # Собираем полный текст: заголовок + содержимое каждого чанка
        text_parts = []
        for chunk in chunks:
            if chunk["title"]:
                text_parts.append(f"## {chunk['title']}\n{chunk['content']}")
            else:
                text_parts.append(chunk["content"])

        # Если чанков нет (документ без заголовков) — берём все параграфы
        if not text_parts:
            all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_parts = all_paragraphs
            chunks = [{"title": "", "content": "\n".join(all_paragraphs), "level": 0}]

        full_text = "\n\n".join(text_parts)

        return ParsedDocument(
            filename=filename,
            content_type=content_type,
            text=full_text,
            pages=[full_text],
            tables=tables,
            chunks=chunks,   # структурные чанки с иерархией — используются RAG
            meta={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(tables),
                "chunk_count": len(chunks),
                "has_headings": any(c["level"] > 0 for c in chunks),
            },
        )

    # XLSX
    # ИСПРАВЛЕНО: wb.sheetnames вызывался после wb.close() — теперь сохраняем
    # количество листов до закрытия книги.
    async def _parse_xlsx(
        self, content: bytes, filename: str, content_type: str
    ) -> ParsedDocument:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheets_text: list[str] = []
        tables: list[list[list[str]]] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[list[str]] = []

            for row in ws.iter_rows(values_only=True):
                row_values = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in row_values):
                    rows.append(row_values)

            if rows:
                tables.append(rows)
                sheet_text = f"=== Лист: {sheet_name} ===\n"
                sheet_text += "\n".join("\t".join(r) for r in rows)
                sheets_text.append(sheet_text)

        # ВАЖНО: сохраняем количество листов ДО закрытия
        sheet_count = len(wb.sheetnames)
        wb.close()

        full_text = "\n\n".join(sheets_text)

        return ParsedDocument(
            filename=filename,
            content_type=content_type,
            text=full_text,
            pages=sheets_text,
            tables=tables,
            meta={"sheet_count": sheet_count},
        )

    # TXT
    # ИСПРАВЛЕНО: переменная encoding могла быть не определена в блоке for/else.
    async def _parse_txt(
        self, content: bytes, filename: str, content_type: str
    ) -> ParsedDocument:
        encoding = "utf-8"  # дефолт на случай если все попытки провалятся

        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                text = content.decode(enc)
                encoding = enc
                break
            except UnicodeDecodeError:
                continue
        else:
            # Все кодировки провалились — декодируем с заменой
            text = content.decode("utf-8", errors="replace")

        return ParsedDocument(
            filename=filename,
            content_type=content_type,
            text=text,
            pages=[text],
            meta={"encoding": encoding, "char_count": len(text)},
        )

    # Изображения (OCR)
    async def _parse_image(
        self, content: bytes, filename: str, content_type: str
    ) -> ParsedDocument:
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image, lang="rus+eng")

            return ParsedDocument(
                filename=filename,
                content_type=content_type,
                text=text.strip(),
                pages=[text.strip()],
                meta={
                    "width": image.width,
                    "height": image.height,
                    "mode": image.mode,
                    "ocr": True,
                },
            )
        except ImportError:
            return ParsedDocument(
                filename=filename,
                content_type=content_type,
                text="",
                error="OCR недоступен: pytesseract не установлен",
            )


# Синглтон
_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser