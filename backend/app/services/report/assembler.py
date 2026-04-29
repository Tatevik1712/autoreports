"""
Сборка итогового DOCX-отчёта из сгенерированных секций.
Ошибки нормоконтроля встраиваются прямо в документ (красным текстом).
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.core.logging import get_logger
from app.schemas.schemas import TemplateSchema, ValidationErrorItem

logger = get_logger(__name__)

# Цвета
COLOR_ERROR = RGBColor(0xC0, 0x00, 0x00)    # тёмно-красный
COLOR_WARNING = RGBColor(0xFF, 0x8C, 0x00)   # оранжевый
COLOR_SECTION = RGBColor(0x1F, 0x49, 0x7D)   # тёмно-синий (заголовки)


class DocxAssembler:
    """Собирает DOCX из словаря section_id → text."""

    def build(
        self,
        template_schema: TemplateSchema,
        sections: dict[str, str],
        validation_errors: list[ValidationErrorItem],
        report_title: str,
    ) -> bytes:
        """Возвращает байты DOCX-файла."""
        doc = Document()
        self._setup_styles(doc)

        # ── Титульный блок ───────────────────────────────────────────────
        title_para = doc.add_heading(report_title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_type_para = doc.add_paragraph(template_schema.document_type)
        doc_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_type_para.runs[0].font.size = Pt(12)

        doc.add_paragraph()  # пустая строка

        # ── Сводка ошибок (если есть) ────────────────────────────────────
        errors_by_section: dict[str | None, list[ValidationErrorItem]] = {}
        for err in validation_errors:
            errors_by_section.setdefault(err.section_id, []).append(err)

        if validation_errors:
            self._add_error_summary(doc, validation_errors)

        # ── Секции ───────────────────────────────────────────────────────
        for section_schema in template_schema.sections:
            sid = section_schema.id
            section_text = sections.get(sid, "")
            section_errors = errors_by_section.get(sid, [])

            # Заголовок секции
            heading = doc.add_heading(section_schema.title, level=1)
            heading.runs[0].font.color.rgb = COLOR_SECTION

            # Текст секции
            if section_text:
                para = doc.add_paragraph(section_text)
                para.style.font.size = Pt(12)
            else:
                self._add_error_inline(
                    doc,
                    "[ДАННЫЕ ОТСУТСТВУЮТ — секция не заполнена]",
                    severity="error",
                )

            # Встроенные ошибки нормоконтроля
            for err in section_errors:
                self._add_error_inline(doc, f"⚠ {err.message} → {err.recommendation}", err.severity)

        # ── Метаданные в конце ───────────────────────────────────────────
        doc.add_page_break()
        self._add_metadata_section(doc, validation_errors)

        # ── Сериализация ─────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        logger.info(
            "docx_assembled",
            title=report_title,
            sections=len(sections),
            errors=len(validation_errors),
            size_bytes=buf.getbuffer().nbytes,
        )

        return buf.read()

    def _setup_styles(self, doc: Document) -> None:
        """Настройка базовых стилей документа."""
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

    def _add_error_summary(
        self, doc: Document, errors: list[ValidationErrorItem]
    ) -> None:
        """Блок в начале документа со сводкой всех ошибок."""
        doc.add_heading("⚠ Сводка замечаний нормоконтроля", level=2)

        error_count = sum(1 for e in errors if e.severity == "error")
        warning_count = sum(1 for e in errors if e.severity == "warning")

        summary = doc.add_paragraph()
        run = summary.add_run(
            f"Найдено ошибок: {error_count}, предупреждений: {warning_count}. "
            "Отчёт сформирован, но требует доработки."
        )
        run.font.color.rgb = COLOR_ERROR
        run.font.bold = True

        for _i, err in enumerate(errors, 1):
            p = doc.add_paragraph(style="List Number")
            color = COLOR_ERROR if err.severity == "error" else COLOR_WARNING
            run = p.add_run(f"[{err.type}] {err.message}")
            run.font.color.rgb = color
            rec_run = p.add_run(f"\n  → Рекомендация: {err.recommendation}")
            rec_run.font.italic = True

        doc.add_paragraph()

    def _add_error_inline(
        self, doc: Document, message: str, severity: str = "warning"
    ) -> None:
        """Добавляет строку с ошибкой прямо внутри документа."""
        p = doc.add_paragraph()
        color = COLOR_ERROR if severity == "error" else COLOR_WARNING
        run = p.add_run(message)
        run.font.color.rgb = color
        run.font.italic = True
        run.font.size = Pt(10)

    def _add_metadata_section(
        self, doc: Document, errors: list[ValidationErrorItem]
    ) -> None:
        """Технические метаданные в конце документа."""
        doc.add_heading("Служебная информация", level=2)
        meta_lines = [
            f"Всего замечаний: {len(errors)}",
            f"Критичных ошибок: {sum(1 for e in errors if e.severity == 'error')}",
            f"Предупреждений: {sum(1 for e in errors if e.severity == 'warning')}",
        ]
        for line in meta_lines:
            p = doc.add_paragraph(line, style="List Bullet")
            p.runs[0].font.size = Pt(10)


def get_docx_assembler() -> DocxAssembler:
    return DocxAssembler()
